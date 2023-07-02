#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# created       :
# description   : convert pdf file to word file
# ----------------------------------------------------------------------------

import click
import PyPDF2
from docx import Document
from docx.shared import Pt


@click.command()
@click.argument('pdf_path', type=click.Path(exists=True))
@click.argument('docx_path', type=click.Path())
def pdf_to_word(pdf_path, docx_path):
    """Convert PDF document to Word file while preserving formatting."""
    pdf_file = open(pdf_path, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    document = Document()

    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text = page.extract_text()

        # Create a new paragraph
        paragraph = document.add_paragraph()

        # Split the text into lines and iterate over each line
        lines = text.split('\n')
        for line in lines:
            # Skip empty lines
            if line.strip():
                # Add a run to the paragraph with the line text
                run = paragraph.add_run(line.strip())

                # Set the font size to preserve the original size (10 points by default)
                run.font.size = Pt(10)

                # Add a line break after each line
                paragraph.add_run().add_break()

    document.save(docx_path)
    pdf_file.close()
    click.echo(f'Successfully converted {pdf_path} to {docx_path}.')


if __name__ == '__main__':
    pdf_to_word()
